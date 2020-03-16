from app import app, db, stats, round_up, round_down
from app.models import *
from werkzeug.utils import secure_filename
from datetime import datetime
import os
import json
from sqlalchemy import and_
from docx import Document

settings = {}
borrel_mode_enabled = False
borrel_mode_drinks = []
overview_emails = False
debt_emails = False

biertje_kwartiertje_participants = []  # Same as an order object: {'user_id': int, 'amount': int}
biertje_kwartiertje_time = 15
biertje_kwartiertje_drink = -1


def fix_float_errors_in_user_balances():
    users = User.query.all()
    for i in range(0, len(users)):
        users[i].balance = round(users[i].balance, 2)
    db.session.commit()


def fix_float_errors_in_usergroup_profits():
    usergroups = Usergroup.query.all()
    for i in range(0, len(usergroups)):
        usergroups[i].profit = round(usergroups[i].profit, 2)
    db.session.commit()


# Because 1.1 + 2.2 !+ 3.3 in Python, we have to fix this every now and then (so at startup seems like a good time)
fix_float_errors_in_user_balances()
fix_float_errors_in_usergroup_profits()


def initialize_settings():
    global settings
    # All settings and their default values
    default_settings = {'dinner_product_id': '-1',
                        'beer_product_id': '-1',
                        'flugel_product_id': '-1',
                        'borrel_mode_user': '-1',
                        'borrel_mode_drinks': '[]',
                        'borrel_mode_amount': '0',
                        'borrel_mode_start_amount': '0',
                        'last_overview_email': '2019-07-01',
                        'last_debt_email': '2019-07-01',
                        'treasurer-email': "",
                        'birthday_groups': "[]"}
    # Get all settings that are in the database
    sses = Setting.query.all()
    settings_keys = []
    # Add all setting keys to the list
    for s in sses:
        settings_keys.append(s.key)

    # Loop over all default settings
    for k, v in default_settings.items():
        # If the default setting is not in the database...
        if k not in settings_keys:
            # Add it with its default value
            s = Setting(key=k, value=v)
            db.session.add(s)
            db.session.commit()

    # Finally, add all settings to the settings object
    # This includes duplicates for the stats (see stats.py)
    for s in Setting.query.all():
        settings[s.key] = s.value


# Run initialization of settings
initialize_settings()

# Determine whether an option should be given to send emails with monthly overviews
now = datetime.now()
last_overview_email = datetime.strptime(settings.get('last_overview_email'), "%Y-%m-%d")
last_debt_email = datetime.strptime(settings.get('last_debt_email'), "%Y-%m-%d")
if now > last_overview_email and now.month != last_overview_email.month and now.day > 3:
    overview_emails = True
if now > last_debt_email and (now - last_debt_email).days >= 3:
    debt_emails = True


@app.context_processor
def emails():
    return dict(send_overview_emails=overview_emails, send_debt_emails=debt_emails)


def update_settings(key, value):
    # Change the settings entry accordingly
    settings[key] = value
    # Get the database entry
    s = Setting.query.get(key)
    # Change the value
    s.value = value
    # Save changes in database
    db.session.commit()


def borrel_mode(drink_id=None):
    global borrel_mode_enabled, borrel_mode_drinks

    # Get the user that is paying for everything in borrel mode
    borrel_mode_user = int(settings['borrel_mode_user'])
    # If the user is -1, borrel mode is disabled so we do not have to execute anything
    if borrel_mode_enabled:
        # If the current drink is in the products...
        if drink_id is None or drink_id in borrel_mode_drinks:
            total_bought = 0
            # We count the total amount of borrel mode products bought (also before the borrel started)
            for p1 in borrel_mode_drinks:
                for p in Purchase.query.filter(Purchase.user_id == borrel_mode_user, Purchase.product_id == p1,
                                               Purchase.price > 0).all():
                    total_bought += p.amount
            total_bought = round_up(total_bought)

            # Calculate how many products are left
            left_over = round_up(float(settings['borrel_mode_amount']), 2) + round_up(float(settings['borrel_mode_start_amount']), 2) - total_bought
            # If it is zero or less...
            if left_over <= 0:
                # We disable borrel mode by changing some settings and variables
                update_settings('borrel_mode_user', -1)
                borrel_mode_enabled = False
                return None

            # If borrel mode is on, we set the boolean and return how many is left and who is paying
            borrel_mode_enabled = True
            return {'left': round_down(left_over, 0),
                    'user': User.query.get(borrel_mode_user).name}

    return None


def set_borrel_mode(products, user_id, amount):
    global borrel_mode_drinks, borrel_mode_enabled

    # The products variable is a list of strings, so we convert each string to an integer and add it to the list
    borrel_mode_drinks = []
    for p in products:
        borrel_mode_drinks.append(int(p))

    total_bought = 0
    # We count the total amount of borrel mode products bought (also before the borrel started)
    for p1 in products:
        for p in Purchase.query.filter(Purchase.user_id == user_id, Purchase.product_id == p1,
                                       Purchase.price > 0).all():
            total_bought += p.amount
    total_bought = round_up(total_bought)

    # Update the database
    update_settings('borrel_mode_start_amount', str(total_bought))
    update_settings('borrel_mode_amount', str(amount))
    update_settings('borrel_mode_user', str(user_id))
    update_settings('borrel_mode_drinks', str(borrel_mode_drinks))

    borrel_mode_enabled = True


def remove_existing_file(filename):
    # Get the location of the file
    file_loc = os.path.join(app.config["UPLOAD_FOLDER"] + filename)
    # If it exists, delete it
    if os.path.exists(file_loc):
        os.remove(file_loc)


def create_filename(product, old_filename, image, character):
    # Get the filename and its extension
    filename, file_extension = os.path.splitext(secure_filename(image.filename))
    # Create a new filename with the productID, a letter, a counter (which we increment by one) and the extension
    filename = str(product.id) + character + "-" + str(int(old_filename.split('-')[1]) + 1) + file_extension
    # Delete the file if it already exists to prevent duplicates
    remove_existing_file(filename)
    # Return the filename
    return filename


def addpurchase(drink_id, user_id, quantity, rondje, price_per_one):
    # Round quantity to at most two decimals
    quantity = round_up(quantity)
    # Get drink and user objects from database
    drink = Product.query.get(drink_id)
    user = User.query.get(user_id)

    # If the price is zero, we do not have to take any inventory
    if price_per_one > 0:
        # If we product is not a mix, we can simply take it from inventory
        if drink.recipe_input is None:
            inventory = take_from_inventory(user, drink_id, quantity)
        # However, if the product is a mix...
        else:
            inventory = {'costs': 0, 'inventory_usage': []}
            # For every ingredient of the mix...
            for r in Recipe.query.filter(Recipe.product_id == drink.id).all():
                # Take the respective quantity from inventory
                result = take_from_inventory(user, r.ingredient_id, round_up(r.quantity * quantity))
                # Add the costs of the ingredient to the total costs
                inventory['costs'] = inventory['costs'] + result['costs']
                # Add the inventory usage
                inventory['inventory_usage'] = inventory['inventory_usage'] + result['inventory_usage']
    else:
        # We take no inventory if the price is zero, so the costs are zero
        inventory = {'costs': 0, 'inventory_usage': []}

    # Get the profitgroup of the user
    profitgroup = Usergroup.query.get(user.profitgroup_id)
    # Calculate the profit
    prof = round_down((price_per_one * quantity - inventory['costs']) * app.config['PROFIT_PERCENTAGE'])
    # Add the profit to the group
    profitgroup.profit = profitgroup.profit + prof
    # Create a row object for the profit table
    profit_obj = Profit(profitgroup_id=profitgroup.id, timestamp=datetime.now(),
                        percentage=app.config['PROFIT_PERCENTAGE'], change=prof, new=profitgroup.profit)
    db.session.add(profit_obj)
    # Save to database
    db.session.commit()

    print(profit_obj.percentage)

    # Calculate the new user balance
    user.balance = user.balance - round_up(float(price_per_one) * quantity)
    # Create a purchase entry in the table, so we can use its purchase ID to create the transaction
    purchase = Purchase(user_id=user.id, timestamp=datetime.now(), product_id=drink.id, price=price_per_one,
                        amount=quantity, round=rondje)
    db.session.add(purchase)
    db.session.commit()

    # Create all inventory usage entries
    for x in inventory['inventory_usage']:
        i_u = Inventory_usage(purchase_id=purchase.id, inventory_id=x['id'], quantity=x['quantity'])
        db.session.add(i_u)
        db.session.commit()
        # Set the object to None to clean up after creation
        i_u = None

    # Calculate the change in balance
    balchange = round_down(-price_per_one * quantity)
    # Create a transaction entry and add it to the database
    transaction = Transaction(user_id=user.id, timestamp=datetime.now(), purchase_id=purchase.id,
                              profit_id=profit_obj.id, balchange=balchange, newbal=user.balance)
    db.session.add(transaction)
    db.session.commit()

    # Update the daily stats with the new purchase
    stats.update_daily_stats_product(drink_id, quantity)
    stats.update_daily_stats('euros', balchange)
    stats.update_daily_stats_drinker(user_id)
    # If the price is zero, we do not add this purchase as it is added somewhere else
    if price_per_one > 0 and drink_id != settings['dinner_product_id']:
        stats.update_daily_stats('purchases', 1)
    if rondje:
        stats.update_daily_stats('rounds', 1)
    return quantity, drink.name, user.name, "success"


def addbalance(user_id, description, amount, profit_id=None):
    # Create an upgrade entry and add it to the database
    upgrade = Upgrade(user_id=user_id, description=description, timestamp=datetime.now(), amount=amount)
    db.session.add(upgrade)
    db.session.commit()
    # Get the user object
    user = User.query.get(upgrade.user_id)
    # Upgrade its balance
    user.balance = user.balance + float(amount)
    # Create a transaction object with the new balance and the upgrade entry and add it to the database
    transaction = Transaction(user_id=user_id, timestamp=datetime.now(), upgrade_id=upgrade.id,
                              balchange=upgrade.amount, profit_id=profit_id, newbal=user.balance)
    db.session.add(transaction)
    db.session.commit()

    # Update the daily stats accordingly
    stats.update_daily_stats('euros', upgrade.amount)

    return upgrade


def add_declaration(user_id, description, amount, group_id):
    # If the group is not 0 aka the bar...
    if group_id is not 0:
        # Get the group and the user for this declaration
        group = Usergroup.query.get(group_id)
        user = User.query.get(user_id)
        # Decrease its profit with the declarated amount
        group.profit -= amount
        # Create the (negative) profit record
        profit = Profit(profitgroup_id=group_id, timestamp=datetime.now(), percentage=1.0, change=-amount, new=group.profit,
                        description=description + " voor " + user.name)
        # Add it to the database
        db.session.add(profit)
        db.session.commit()

        # Get the ID for when the balance gets added
        profit_id = profit.id
    else:
        # No profit has been removed, so the profit_id does not exist
        profit_id = None

    # Add the balance
    return addbalance(user_id, description, amount, profit_id)



def parse_recipe(recipe):
    # If there is a recipe, let's parse it!
    if recipe != "":
        # The format is "1x1, 2x1, ..." where the first digit is the quantity and the second digit is the productID
        # First, split the recipe into components (products)
        components = recipe.split(",")
        result_recipe = {}
        # For all components...
        for c in components:
            # Delete the spaces from the string and split it on the x (so data[0] is the quantity and data[1] is the
            #     productID
            data = c.replace(" ", "").split("x")
            # Verify the data is in the corredt format
            if int(data[0]) <= 0 or int(data[1]) <= 0 or len(data) != 2:
                return "Recept voldoet niet aan de gegeven syntax!", "danger"
            # Check that there are no other symbols in the components objects
            len_a = int(len(data[0])) + int(len(data[1])) + 1
            len_b = int(len(c.replace(" ", "")))
            if len_a != len_b:
                return "Recept voldoet niet aan de gegeven syntax!", "danger"
            # Change the strings to integers
            data[0] = int(data[0])
            data[1] = int(data[1])
            # Ingredient must exist and has to be purchaseable
            if Product.query.get(data[1]) is None or Product.query.get(data[1]).purchaseable is False:
                return "Product met ID {} bestaat niet of is niet beschikbaar!".format(str(data[1])), "danger"
            # Ingredient cannot be a mix itself
            if Recipe.query.filter(Recipe.product_id == data[1]).count() > 0:
                return "Ingrediënt met ID {} bestaat zelf ook uit ingredienten!".format(str(data[1])), "danger"
            # Add the ingedient to the result recipe with its amount
            result_recipe[data[1]] = data[0]
        # If we have only one ingredient, we return an error
        if len(result_recipe) <= 1:
            return "Een recept kan niet uit één type product bestaan!", "danger"
        # If all is right, we return the resulting (parsed) recipe
        return result_recipe
    return None


def adddrink(name, price, category, order, image, hoverimage, recipe, inventory_warning, alcohol, volume, unit):
    # Get the recipe if it exists
    result_recipe = parse_recipe(recipe)
    # If the result_recipe is a tuple, it is a return message so we return it
    if type(result_recipe) is tuple:
        return result_recipe
    # If no inventory warning is set, we set it to zero by default
    if inventory_warning is None:
        inventory_warning = 0

    # Get the filename of both images and check that they have the correct extensions
    s_filename, s_file_extension = os.path.splitext(secure_filename(image.filename))
    if s_file_extension[1:] not in app.config["ALLOWED_EXTENSIONS"]:
        return "Statische afbeelding is niet van het correcte bestandstype (geen .png, .jpg, .bmp of .gif)", "danger"
    if hoverimage != "":
        h_filename, h_file_extension = os.path.splitext(secure_filename(hoverimage.filename))
        if h_file_extension[1:] not in app.config["ALLOWED_EXTENSIONS"]:
            return "Hover afbeelding is niet van het correcte bestandstype (geen .png, .jpg, .bmp of .gif)", "danger"

    # Calculate the product placement in the order list
    amount_of_prod = Product.query.count()
    # If the order number is smaller than the amount...
    if order < amount_of_prod:
        # We need to increase the order number of all successive products
        p_ord = Product.query.filter(Product.order >= order).all()
        for i in range(0, len(p_ord)):
            p_ord[i].order = p_ord[i].order + 1
    # If this is not the case, we set the position to the last product in the list
    elif order >= amount_of_prod:
        order = amount_of_prod + 1

    # If we have a recipe...
    if result_recipe is not None:
        # We create a product with the recipe
        product = Product(name=name, price=price, category=category, order=order, purchaseable=True, image="",
                          hoverimage="",
                          recipe_input=result_recipe)
    # If we do not have a recipe...
    else:
        # We set the default values of volume and alcohol if none is given
        if volume is "":
            volume = "0"
        if alcohol is "":
            alcohol = "0"
        # Then, we create a new product in the database
        product = Product(name=name, price=price, category=category, order=order, purchaseable=True, image="",
                          hoverimage="",
                          inventory_warning=inventory_warning, volume=int(float(volume)), unit=unit,
                          alcohol=float(alcohol.replace(",", ".").replace("%", "").replace(" ", "")) / 100)
    db.session.add(product)
    db.session.commit()

    # If we have a recipe...
    if result_recipe is not None:
        # Then we need to add a recipe entry for all ingredients in the database
        for key, val in result_recipe.items():
            db.session.add(Recipe(product_id=product.id, ingredient_id=key, quantity=val))
        db.session.commit()

    # Create the upload folder if it does not exist
    if not os.path.exists(app.config["UPLOAD_FOLDER"]):
        os.makedirs(app.config["UPLOAD_FOLDER"])
    # Set the product images
    product.image = create_filename(product, str(product.id) + "s-0", image, "s")
    image.save(os.path.join(app.config["UPLOAD_FOLDER"], product.image))

    if hoverimage == "":
        product.hoverimage = product.image
    else:
        product.hoverimage = create_filename(product, str(product.id) + "h-0", hoverimage, "h")
        hoverimage.save(os.path.join(app.config["UPLOAD_FOLDER"], product.hoverimage))
    db.session.commit()

    # Update the daily stats because we added a product
    stats.update_daily_stats('products', 1)

    return "Product {} succesvol aangemaakt".format(product.name), "success"


def adduser(name, email, group, profitgroup, birthday):
    user = User(name=name, email=email, usergroup_id=group, profitgroup_id=profitgroup, birthday=birthday)
    db.session.add(user)
    db.session.commit()

    stats.update_daily_stats('users', 1)

    return "Gebruiker {} succesvol geregistreerd".format(user.name), "success"


def addusergroup(name):
    usergroup = Usergroup(name=name, )
    db.session.add(usergroup)
    db.session.commit()
    return "Groep {} succesvol aangemaakt".format(usergroup.name), "success"


def addquote(q):
    quote = Quote(timestamp=datetime.now(), value=q)
    db.session.add(quote)
    db.session.commit()


def add_sound(name, key, code, file):
    sound = Sound(name=name, keyboard_key=key, keyboard_code=code)
    db.session.add(sound)
    db.session.commit()

    filename, extension = os.path.splitext(secure_filename(file.filename))
    sound.filename = str(sound.id) + "_" + filename + extension
    file.save(os.path.join(app.config['SOUNDBOARD_FOLDER'], sound.filename))
    db.session.commit()


def deluser(user_id):
    user = User.query.get(user_id)
    name = user.name
    for t in user.transactions.all():
        db.session.delete(t)
    for u in user.upgrades.all():
        db.session.delete(u)
    negative_inventory_found = False
    for p in user.purchases.all():
        for i in Inventory_usage.query.filter(Inventory_usage.purchase_id == p.id).all():
            if i.amount < 0:
                negative_inventory_found = True
        db.session.delete(p)

    if negative_inventory_found:
        db.session.rollback()
        return "Verwijderen van gebruiker {} mislukt: er staat nog negatieve inventaris op zijn naam!".format(
            name), "danger"
    else:
        db.session.delete(user)
        db.session.commit()

        stats.update_daily_stats('users', -1)

        return "Gebruiker {} verwijderd".format(name), "success"


def del_profit(profit_id):
    # Find the profit object
    profit = Profit.query.get(profit_id)
    # Get the profit group and undo the added profit
    profitgroup = Usergroup.query.get(profit.profitgroup_id)
    profitgroup.profit -= profit.change

    # For all newer recorded profits for the same group...
    for p in Profit.query.filter(Profit.profitgroup_id == profit.profitgroup_id, Profit.timestamp > profit.timestamp).all():
        # Reduce its new profit balance with the value of the to-be removed row
        p.new = p.new - profit.change
    # Actually delete the row
    db.session.delete(profit)
    db.session.commit()


def deltransaction(transaction_id):
    # Get the tranaction and its corresponding user and profitgroup
    transaction = Transaction.query.get(transaction_id)
    user = User.query.get(transaction.user_id)
    profitgroup = Usergroup.query.get(user.profitgroup_id)

    # If the transaction is an upgrade...
    if transaction.purchase_id is None:

        # If this upgrade was a declaration to a user group, we have to undo this declaration
        if transaction.profit_id is not None:
            # Delete the profit object
            del_profit(transaction.profit_id)

        # Get the upgrade object and delete it!
        # The balance will be changed later
        upgrade = Upgrade.query.get(transaction.upgrade_id)
        db.session.delete(upgrade)
    else:
        purchase = Purchase.query.get(transaction.purchase_id)

        # If the transaction is from before version 1.4.4.1 or earlier...
        if transaction.profit_id is None:
            # Remove the profit the old way
            # This is the first step, that removes too much profit
            profitgroup.profit = profitgroup.profit - purchase.price * purchase.amount
        else:
            # If this is not the case, delete the profit object
            del_profit(transaction.profit_id)

        inv_usage = Inventory_usage.query.filter(Inventory_usage.purchase_id == purchase.id).all()
        if type(inv_usage) is not list:
            inv_usage = [inv_usage]
        for i in range(0, len(inv_usage)):
            inventory = Inventory.query.get(inv_usage[i].inventory_id)
            inventory.quantity = inventory.quantity + inv_usage[i].quantity
            # If the transaction is from before version 1.4.4.1 or earlier...
            if transaction.profit_id is None:
                # Do the second step: increase the profit again with the price before profit
                # (so only the profit has been removed)
                profitgroup.profit = profitgroup.profit + inventory.price_before_profit * inv_usage[i].quantity
            db.session.delete(inv_usage[i])
            db.session.commit()
            inventory = None

        recipe = Recipe.query.filter(Recipe.product_id == purchase.product_id).all()
        if len(recipe) == 0:
            fix_negative_inventory(purchase.product_id)
        else:
            for r in recipe:
                fix_negative_inventory(r.ingredient_id)
        db.session.delete(purchase)

    for t in Transaction.query.filter(Transaction.user_id == transaction.user_id,
                                      Transaction.timestamp > transaction.timestamp).all():
        t.newbal = t.newbal - transaction.balchange
    user.balance = user.balance - transaction.balchange
    db.session.delete(transaction)
    db.session.commit()
    return "Transactie met ID {} succesvol verwijderd!".format(transaction.id), "success"


def deldrink(drink_id):
    product = Product.query.get(drink_id)
    product.purchaseable = False
    db.session.commit()
    return 'Product {} is niet meer beschikbaar'.format(product.name), "success"


def delusergroup(usergroup_id):
    usergroup = Usergroup.query.get(usergroup_id)
    db.session.delete(usergroup)
    db.session.commit()
    return "Groep {} verwijderd".format(usergroup.name), "success"


def del_sound(sound_id):
    sound = Sound.query.get(sound_id)
    db.session.delete(sound)
    db.session.commit()
    return "Geluid {} verwijderd".format(sound.name), "success"


def edit_profit(profit, amount):
    # Calculate the new profit while taking the percentage for the bar into account
    old_change = profit.change
    old_change_before_perc = old_change / profit.percentage
    new_change_before_perc = old_change_before_perc + amount
    profit.change = new_change_before_perc * profit.percentage
    diff = profit.change - old_change
    # For all newer recorded profits for the same group...
    for p in Profit.query.filter(Profit.profitgroup_id == profit.profitgroup_id, Profit.timestamp > profit.timestamp).all():
        # Reduce its new profit balance with the amount
        p.new += diff
    
    # Get the profit group
    profitgroup = Usergroup.query.get(profit.profitgroup_id)
    # Change the profit accordingly
    profitgroup.profit += diff

    # Write all changes
    db.session.commit()


def editdrink_attr(product_id, name, price, category, order, purchaseable, recipe, inventory_warning, alcohol, volume,
                   unit):
    result_recipe = parse_recipe(recipe)
    if type(result_recipe) is tuple:
        return result_recipe

    product = Product.query.get(product_id)

    if product.recipe_input != result_recipe:
        for r in Recipe.query.filter(Recipe.product_id == product.id):
            db.session.delete(r)
        for key, val in result_recipe:
            db.session.add(Recipe(product_id=product.id, ingredient_id=key, quantity=val))

    purchaseable_old = product.purchaseable

    product.name = name
    product.price = price
    product.category = category
    if order > Product.query.count():
        order = Product.query.count()
    if order < product.order:
        p_ord = Product.query.filter(Product.order >= order, Product.order < product.order).all()
        for i in range(0, len(p_ord)):
            p_ord[i].order = p_ord[i].order + 1
    elif order > product.order:
        p_ord = Product.query.filter(Product.order <= order, Product.order > product.order).all()
        for i in range(0, len(p_ord)):
            p_ord[i].order = p_ord[i].order - 1
    product.order = order
    product.purchaseable = purchaseable
    product.recipe_input = result_recipe
    product.inventory_warning = inventory_warning
    if alcohol != "":
        product.alcohol = float(str(alcohol).replace(",", ".").replace("%", "").replace(" ", "")) / 100
    else:
        product.alcohol = None
    product.volume = int(float(volume))
    product.unit = unit
    db.session.commit()

    if purchaseable_old is True and purchaseable is False:
        stats.update_daily_stats('products', -1)
    elif purchaseable_old is False and purchaseable is True:
        stats.update_daily_stats('products', 1)

    return "Product {} (ID: {}) succesvol aangepast!".format(product.name, product.id), "success"


def editdrink_price(product_id, price):
    product = Product.query.get(product_id)
    product.price = price
    db.session.commit()


def editdrink_image(product_id, image, hoverimage):
    product = Product.query.get(product_id)
    if image != "":
        s_filename, s_file_extension = os.path.splitext(secure_filename(image.filename))
        if s_file_extension[1:] not in app.config["ALLOWED_EXTENSIONS"]:
            return "Statische afbeelding is niet van het correcte bestandstype (geen .png, .jpg, .bmp of .gif)", "danger"
        product.image = create_filename(product, os.path.splitext(product.image)[0], image, "s")
        image.save(os.path.join(app.config["UPLOAD_FOLDER"], product.image))
    if hoverimage != "":
        h_filename, h_file_extension = os.path.splitext(secure_filename(hoverimage.filename))
        if h_file_extension[1:] not in app.config["ALLOWED_EXTENSIONS"]:
            return "Hover afbeelding is niet van het correcte bestandstype (geen .png, .jpg, .bmp of .gif)", "danger"
        product.hoverimage = create_filename(product, os.path.splitext(product.image)[0], hoverimage, "h")
        hoverimage.save(os.path.join(app.config["UPLOAD_FOLDER"], product.hoverimage))
    db.session.commit()
    return "Afbeelding van product {} (ID: {}) succesvol aangepast!".format(product.name, product.id), "success"


# -- Inventory management -- #

def get_inventory_stock(product_id):
    if Product.query.get(product_id).recipe_input is None:
        return [get_inventory_stock_for_product(product_id)]
    else:
        result = []
        for key, value in Product.query.get(product_id).recipe_input.items():
            p = Product.query.get(key)
            stock = get_inventory_stock_for_product(key)
            stock["name"] = p.name
            stock["recipe_quantity"] = value
            result.append(stock)
        return result


def get_inventory_stock_for_product(product_id):
    result = {}
    inventories = Inventory.query.filter(Inventory.product_id == product_id).all()
    result['inventory_warning'] = Product.query.get(product_id).inventory_warning
    if len(inventories) == 0:
        result['quantity'] = 0
        result['entries'] = 0
        result['oldest'] = None
        result['newest'] = None
    else:
        result['quantity'] = 0
        result['entries'] = len(inventories)
        result['oldest'] = inventories[0].timestamp
        result['newest'] = inventories[-1].timestamp
        for i in inventories:
            result['quantity'] = result['quantity'] + int(i.quantity)
        if result['quantity'] < 0:
            result['oldest'] = None
            result['newest'] = None
    return result


def get_product_stats(product_id):
    result = {}

    if Recipe.query.filter(Recipe.product_id == product_id).count() == 0:
        p = Product.query.get(product_id)
        result['alcohol'] = p.volume * p.alcohol
        result['volume'] = p.volume / 1000
        result['stock'] = [
            {"product": p.name, "quantity": get_inventory(p.id), "inventory_warning": p.inventory_warning}]
    else:
        result['alcohol'] = 0
        result['stock'] = []
        result['volume'] = 0.0
        for r in Recipe.query.filter(Recipe.product_id == product_id).all():
            p = Product.query.get(r.ingredient_id)
            result['alcohol'] = result['alcohol'] + p.volume * p.alcohol
            result['stock'].append(
                {"product": p.name, "quantity": get_inventory(p.id), "inventory_warning": p.inventory_warning})
            result['volume'] = result['volume'] + p.volume / 1000

    purchases = Purchase.query.filter(Purchase.product_id == product_id, Purchase.price > 0).all()
    users = {u.id: [0, 0] for u in User.query.all()}
    result['total_bought'] = 0
    for pur in purchases:
        result['total_bought'] = result['total_bought'] + pur.amount
        if not pur.round:
            users[pur.user_id][0] = users[pur.user_id][0] + pur.amount
        else:
            users[pur.user_id][1] = users[pur.user_id][1] + 1
    largest_consumer = None
    largest_round_giver = None
    largest_consumer_amount = 0
    largest_round_giver_amount = 0
    for user, amount in users.items():
        if amount[0] > largest_consumer_amount:
            largest_consumer = user
            largest_consumer_amount = amount[0]
        if amount[1] > largest_round_giver_amount:
            largest_round_giver = user
            largest_round_giver_amount = amount[1]

    result['largest_consumer'] = {'user': largest_consumer, 'amount': largest_consumer_amount}
    result['largest_rounder'] = {'user': largest_round_giver, 'amount': largest_round_giver_amount}
    return result


def find_inventory(product_id, index):
    inventories = Inventory.query.filter(Inventory.product_id == product_id, Inventory.quantity != 0).all()
    if len(inventories) == 0:
        return None
    elif type(inventories) is Inventory:
        return inventories
    else:
        return inventories[index]


def find_newest_inventory(product_id):
    return find_inventory(product_id, -1)


def find_newest_product_price(product_id):
    inventories = Inventory.query.filter(Inventory.product_id == product_id, Inventory.quantity > 0).all()
    if len(inventories) == 0:
        inventories = Inventory.query.filter(Inventory.product_id == product_id, Inventory.quantity == 0).all()
        if len(inventories) == 0:
            return 0.0
        elif type(inventories) is Inventory:
            return inventories.price_before_profit
        else:
            return inventories[-1].price_before_profit
    elif type(inventories) is Inventory:
        return inventories.price_before_profit
    else:
        return inventories[-1].price_before_profit


def find_oldest_inventory(product_id):
    return find_inventory(product_id, 0)


def get_inventory(product_id):
    inventories = Inventory.query.filter(and_(Inventory.product_id == product_id, Inventory.quantity != 0)).all()
    sum = 0
    for i in inventories:
        sum = sum + i.quantity
    return sum


def add_inventory(p_id, quantity, price_before_profit, note):
    old_inv = find_oldest_inventory(p_id)
    neg_inv = Inventory.query.filter(Inventory.product_id == p_id, Inventory.quantity < 0).all()
    new_inv = Inventory(product_id=p_id, timestamp=datetime.now(), quantity=float(quantity),
                        price_before_profit=price_before_profit, note=note)
    db.session.add(new_inv)
    db.session.commit()
    for n in neg_inv:
        fix_neg_inv(n, new_inv)
        if new_inv.quantity == 0:
            break

    product = Product.query.get(p_id)
    return "{} {} toegevoegd aan inventaris!".format(str(quantity), product.name), "success"


def fix_neg_inv(old_inv, new_inv):
    inv_use = Inventory_usage.query.filter(Inventory_usage.inventory_id == old_inv.id).all()
    for i in range(0, len(inv_use)):
        if new_inv.quantity == 0:
            break
        pur = Purchase.query.get(inv_use[i].purchase_id)
        if pur is not None:
            t = pur.transactions.all()[0]
            user = User.query.get(pur.user_id)
            if user is not None:
                profitgroup = Usergroup.query.get(user.profitgroup_id)
            # If the purchase is created in version 1.5.0.0 or newer, we get the profit object as well
            if t.profit_id is not None:
                profit = Profit.query.get(pur.profit_id)
            else:
                profit = None
        else:
            profitgroup = None
            app.logger.info("There is no purchase attached to an inventory usage row of inventory {}".format(old_inv.id))
            raise ValueError

        if inv_use[i].quantity > new_inv.quantity:
            new_inv_use = Inventory_usage(purchase_id=inv_use[i].purchase_id, inventory_id=new_inv.id,
                                          quantity=new_inv.quantity)
            db.session.add(new_inv_use)
            inv_use[i].quantity = inv_use[i].quantity - new_inv.quantity
            if profitgroup is not None:
                change = -(new_inv.price_before_profit * new_inv_use.quantity)
                # Use the old method if the transaction is from 1.4.4.1 or older
                if profit is None:
                    profitgroup.profit = profitgroup.profit + change
                else:
                    # Use the new method otherwise
                    edit_profit(change)
            new_inv.quantity = 0
            old_inv.quantity = old_inv.quantity + new_inv_use.quantity
            db.session.commit()
            break
        else:
            inv_use[i].inventory_id = new_inv.id
            if profitgroup is not None:
                change = -(new_inv.price_before_profit * inv_use[i].quantity)
                if profit is None:
                    profitgroup.profit = profitgroup.profit + change
                else:
                    edit_profit(change)
            new_inv.quantity = new_inv.quantity - inv_use[i].quantity
            old_inv.quantity = old_inv.quantity + inv_use[i].quantity
            db.session.commit()


def take_from_inventory(user, product_id, quantity):
    product = Product.query.get(product_id)
    if user is not None:
        profitgroup = Usergroup.query.get(user.profitgroup_id)
    else:
        profitgroup = None
    added_costs = 0
    inventory_usage = []
    while quantity > 0:
        inventory = find_oldest_inventory(product_id)
        if inventory is None:
            inventory = Inventory(product_id=product_id, timestamp=datetime.utcnow(), quantity=float(- quantity),
                                  price_before_profit=0.0, note="Noodinventaris")
            db.session.add(inventory)
            db.session.commit()
            inventory_usage.append({"id": inventory.id, "quantity": quantity})
            db.session.commit()
            break
        else:
            if quantity < inventory.quantity:
                if profitgroup is not None:
                    added_costs = added_costs + (inventory.price_before_profit * quantity)
                inventory_usage.append({"id": inventory.id, "quantity": quantity})
                inventory.quantity = inventory.quantity - quantity
                break
            elif quantity >= inventory.quantity > 0.0:
                if profitgroup is not None:
                    added_costs = added_costs + (inventory.price_before_profit * inventory.quantity)
                inventory_usage.append({"id": inventory.id, "quantity": inventory.quantity})
                quantity = quantity - inventory.quantity  # decrease quantity
                inventory.quantity = 0
            else:
                inventory.quantity = inventory.quantity - quantity
                inventory_usage.append({"id": inventory.id, "quantity": quantity})
                break
        db.session.commit()
    db.session.commit()
    return {"costs": added_costs, "inventory_usage": inventory_usage}


def increase_inventory_quantity(product_id, quantity):
    p = Product.query.get(product_id)
    inventory = find_newest_inventory(product_id)
    if inventory is None:
        inventories = Inventory.query.filter(Inventory.product_id == product_id).all()
        if len(inventories) == 0:
            empty_invs = Inventory.query.filter(Inventory.product_id == product_id, Inventory.quantity < 0).all()
            if len(empty_invs) == 0:
                price = 0.0
            else:
                price = empty_invs[-1].price_before_profit
            inventory = Inventory(product_id=product_id, timestamp=datetime.utcnow(), price_before_profit=price,
                                  quantity=0, note="Inventariscorrectie")
            db.session.add(inventory)
        elif type(inventories) is Inventory:
            inventory = Inventory
        else:
            inventory = inventories[-1]
    inventory.quantity = inventory.quantity + quantity
    db.session.commit()


def fix_negative_inventory(p_id):
    neg_inv = Inventory.query.filter(Inventory.product_id == p_id, Inventory.quantity < 0).all()
    if type(neg_inv) is Inventory:
        neg_inv = [neg_inv]

    for n in neg_inv:
        print("N: " + neg_inv.__repr__())
        pos_inv = Inventory.query.filter(and_(
            Inventory.product_id == p_id,
            Inventory.quantity > 0,
            Inventory.timestamp > n.timestamp
        )).all()
        print("P: " + pos_inv.__repr__())
        if type(pos_inv) is Inventory:
            pos_inv = [pos_inv]
        for p in pos_inv:
            fix_neg_inv(n, p)
            if n.quantity == 0:
                break


def correct_inventory(json):
    # Create a list of all groups that will participate in this inventory correction
    all_groups = set()
    for i in json:
        s = set(i['groups'])
        all_groups.update(s)
    all_groups = list(all_groups)
    per_group_costs = [0] * len(all_groups)
    total_costs = 0

    # Create a document for the overview of this correction
    now = datetime.now()
    document = Document()
    document.add_heading('Inventaris Correctie', 0)

    document.add_paragraph("Op {} om {} is er een inventariscorrectie uitgevoerd. Hierbij zijn de volgende wijzigingen "
                           "doorgevoerd in de inventaris.".format(now.strftime("%Y-%m-%d"), now.strftime("%H:%M:%S")))

    # Table layout: Product name | Inventory in Tikker | Real inventory | Difference | Costs | <per group its costs or
    # profit
    table = document.add_table(rows=1, cols=5 + len(all_groups))
    header_cells = table.rows[0].cells
    header_cells[0].text = "Productnaam"
    header_cells[1].text = "Inv in Tikker"
    header_cells[2].text = "Echte inv"
    header_cells[3].text = "Verschil"
    header_cells[4].text = "Kosten"
    for i in range(0, len(all_groups)):
        header_cells[5 + i].text = "Kosten {}".format(Usergroup.query.get(all_groups[i]).name)

    for i in json:
        # Get the product object
        p = Product.query.get(i['product_id'])
        # Calculate the inventory difference
        tikker_inv = calcStock(p.id)
        diff = int(i['stock']) - int(tikker_inv)

        # Create a new row for the document table
        row_cells = table.add_row().cells
        row_cells[0].text = p.name
        row_cells[1].text = str(tikker_inv)
        row_cells[2].text = str(i['stock'])
        row_cells[3].text = str(diff)
        if diff < 0:
            row_cells[4].text = '-€ %.2f' % round_up(-diff * p.price)
            total_costs += round_up(-diff * p.price)
        elif diff > 0:
            row_cells[4].text = '€ %.2f' % round_down(diff * p.price)
            total_costs += round_down(diff * p.price)

        # If the difference is less than 0, we have lost some inventory, which costs money
        if diff < 0:
            for g_id in i['groups']:
                # Get the group object
                g = Usergroup.query.get(g_id)
                # Calculate the costs for this group
                cost = round_down(diff * p.price / len(i['groups']))
                # Change its profit
                g.profit = g.profit + cost
                profit = Profit(profitgroup_id=g_id, timestamp=datetime.now(), percentage=1.0,
                                change=cost, new=g.profit, description="{} {} inventariscorrectie".format(str(diff), p.name))
                db.session.add(profit)
                db.session.commit()
                # Add it to the row in the table and to the cumulative costs
                index = all_groups.index(g_id)
                per_group_costs[index] += cost
                row_cells[5 + index].text = '-€ %.2f' % -cost
                # Reset the group object
                g = None
            print(take_from_inventory(None, p.id, -diff))
        elif diff > 0:
            new_price = find_newest_product_price(p.id)
            add_inventory(p.id, diff, new_price, "Inventariscorrectie")
            for g_id in i['groups']:
                g = Usergroup.query.get(g_id)
                profit_for_group = round_down(diff * new_price / len(i['groups']))
                g.profit = g.profit + profit_for_group
                profit = Profit(profitgroup_id=g_id, timestamp=datetime.now(), percentage=1.0,
                                change=profit_for_group, new=g.profit,
                                description="{} {} inventariscorrectie".format(str(diff), p.name))
                db.session.add(profit)
                db.session.commit()
                # Add it to the row in the table and to the cumulative costs
                index = all_groups.index(g_id)
                per_group_costs[index] += profit_for_group
                row_cells[5 + index].text = '€ %.2f' % profit_for_group
                # Reset the group object
                g = None

    row_cells = table.add_row().cells
    if total_costs < 0:
        row_cells[4].text = '-€ %.2f' % -total_costs
    else:
        row_cells[4].text = '€ %.2f' % total_costs

    for i in range(0, len(per_group_costs)):
        if per_group_costs[i] < 0:
            row_cells[5 + i].text = '-€ %.2f' % -per_group_costs[i]
        else:
            row_cells[5 + i].text = '€ %.2f' % per_group_costs[i]

    # Try to save the file
    # If it is not possible because the file is locked, try it again with a different filename until it works
    count = 0
    while True:
        filename = os.path.join(app.config['DOCUMENT_FOLDER'], 'inventariscorrectie_{}_{}.docx'
                                .format(now.strftime("%Y%m%d"), count))
        # If the file exists, raise the file name and try again
        if os.path.exists(filename):
            count += 1
            continue
        try:
            # If saving the file is successful, escape the loop and finish
            document.save(filename)
            break
        # If the file is opened, catch the error and try again
        except PermissionError:
            count += 1

    return "Inventaris correctie succesvol doorgevoerd! Het rapport is opgeslagen in {}".format(filename), "success"


def payout_profit(usergroup_id, amount, password):
    if password != app.config['ADMIN_PASSWORD']:
        return "Verificatiecode is ongeldig. Probeer het opnieuw", "danger"

    usergroup = Usergroup.query.get(usergroup_id)
    if amount > usergroup.profit:
        return "Het uit te keren bedrag is groter dan de opgebouwde winst! Kies een lager bedrag", "danger"
    if amount == 0:
        return "Het gekozen bedrag is ongeldig"

    usergroup.profit = usergroup.profit - amount
    profit = Profit(profitgroup_id=usergroup_id, timestamp=datetime.now(), percentage=1.0, change=-amount,
                    new=usergroup.profit)
    db.session.add(profit)
    db.session.commit()
    return "€ {} winst van {} uitgekeerd uit Tikker".format(str('%.2f' % usergroup.profit),
                                                            usergroup.name), "success"


def calcStock(product_id):
    inventories = Inventory.query.filter(and_(Inventory.product_id == product_id, Inventory.quantity != 0))
    sum = 0
    for i in inventories:
        sum = sum + i.quantity
    return sum


def calc_inventory_value(product_id):
    sum = 0
    for i in Inventory.query.filter(and_(Inventory.product_id == product_id, Inventory.quantity != 0)):
        sum += i.quantity * i.price_before_profit
    return sum


def rollback():
    db.session.rollback()


def force_edit():
    users = [1, 3]
    print(User.query.filter(User.id.in_(users)).all())


def is_birthday():
    today = datetime.today()
    birthday_groups = json.loads(settings['birthday_groups'])
    days_in_year = 365.2425
    birthdays = []
    for u in User.query.all():
        bday = datetime(today.year, u.birthday.month, u.birthday.day)
        diff = (today - bday).days
        if (0 <= diff <= 7) and (u.usergroup_id in birthday_groups):
            age = int((today - u.birthday).days / days_in_year)
            birthdays.append({"user": u, "age": age})
    return birthdays


def upcoming_birthdays():
    today = datetime.today()
    birthday_groups = json.loads(settings['birthday_groups'])
    days_in_year = 365.2425
    birthdays = []
    for u in User.query.all():
        if u.usergroup_id in birthday_groups:
            bday = datetime(today.year, u.birthday.month, u.birthday.day)
            new_age = int(round_up((today - u.birthday).days / days_in_year, 0))
            if bday < today:
                bday = datetime(today.year + 1, u.birthday.month, u.birthday.day)
            diff = (bday - today).days + 1
            birthdays.append({'user': u.name,
                              'birthday': bday,
                              'age': new_age,
                              'days': diff})
    return sorted(birthdays, key=lambda i: i['birthday'])
